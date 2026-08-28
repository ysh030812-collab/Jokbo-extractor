# -*- coding: utf-8 -*-
"""회귀 테스트용 DOCX 픽스처를 tests/fixtures/ 에 만든다.

실제 족보(2022·2023 감면 기말)에서 파서가 걸려 넘어졌던 구간만 그대로 옮겼다.
저작물이므로 문항 내용 자체는 넣지 않고, 번호·구분점·선지 구조만 재현한다.
"""
import os, pathlib
import docx
from docx.oxml import parse_xml

OUT = pathlib.Path(__file__).parent / "fixtures"
OUT.mkdir(exist_ok=True)

NUMPR = '''<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
 <w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr></w:pPr>'''

TEXTBOX = '''<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
 xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
 xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
 xmlns:v="urn:schemas-microsoft-com:vml"
 xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
 xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
 <mc:AlternateContent><mc:Choice Requires="wps"><w:drawing><wp:inline><a:graphic><a:graphicData>
   <wps:wsp><wps:txbx><w:txbxContent><w:p><w:r><w:t>TEXTBOX_CONTENT</w:t></w:r></w:p></w:txbxContent></wps:txbx></wps:wsp>
 </a:graphicData></a:graphic></wp:inline></w:drawing></mc:Choice>
 <mc:Fallback><w:pict><v:shape><v:textbox><w:txbxContent><w:p><w:r><w:t>TEXTBOX_CONTENT</w:t></w:r></w:p></w:txbxContent></v:textbox></v:shape></w:pict></mc:Fallback>
 </mc:AlternateContent></w:r>'''


# ── 2022 기말: 사용자가 보고한 143/144 병합을 포함한 실제 형태 ──────────
d = docx.Document()
A = d.add_paragraph
A("2022학년도 감염과 면역 Ⅰ 기말고사")
# ── 1번: 번호가 통째로 빠져 있다 (원본 그대로) ─────────────────────
A("다음 그림은 광견병 예방의 방침을 결정하는 알고리즘을 나타낸 것이다. 옳은 것은?")
A("하나를 선택하세요.")
A("1. 개에게 물린 즉시 백신을 접종한다.")
A("2. 관찰 가능한 개는 10일간 관찰한다.")
A("3. 면역글로불린은 투여하지 않는다.")
A("4. 야생동물은 관찰 후 결정한다.")
A("5. 상처는 씻지 않는다.")
# ── 2번: 마침표 뒤 공백 없음 ────────────────────────────────────
A("2.다음 도표는 백일해 발생률의 연도별 추이이다. 옳지 않은 것은?")
A("하나를 선택하세요.")
A("1. 1980년대 이후 감소하였다.")
A("2. 백신 도입과 관련이 있다.")
A("3. 최근 다시 증가 추세이다.")
A("4. 성인 감염은 드물다.")
A("5. Tdap 추가접종이 권고된다.")
# ── 3번: 선지가 4개뿐 ───────────────────────────────────────────
A("3. 다음 중 인터페론의 작용으로 옳은 것은?")
A("1. 세포 내 항바이러스 상태를 유도한다.")
A("2. 보체를 직접 활성화한다.")
A("3. 세균의 세포벽을 분해한다.")
A("4. 항체를 직접 생산한다.")
# ── 4번: 닫는 괄호 표기 ─────────────────────────────────────────
A("4) 다음 중 옳은 것은?")
A("1. 첫째 선지")
A("2. 둘째 선지")
A("3. 셋째 선지")
A("4. 넷째 선지")
A("5. 다섯째 선지")
# ── 5번: <보기> 가 딸린 문제 ────────────────────────────────────
A("5. 다음 <보기> 중 기생충 질환을 모두 고른 것은?")
A("<보기>")
A("a. 수면병")
A("b. 옴")
A("c. 개조충")
A("1. a, b")
A("2. a, c")
A("3. b, c")
A("4. a, b, c")
A("5. c만")
# ── 106번: 선지가 16개 (실제 2023 87번이 이 형태) ────────────────
A("106. 종양 유발과 가장 연관된 바이러스는?")
for i, t in enumerate(["인플루엔자 1", "인플루엔자 2", "인플루엔자 3", "간염바이러스 A",
                       "간염바이러스 B", "간염바이러스 C", "간염바이러스 D", "간염바이러스 E",
                       "HPV 1,2,3,4", "HPV 6,11", "HPV 16,18", "아데노 1,2,5",
                       "아데노 3,7", "아데노 8,19,37", "아데노 40,41", "아데노 11,21"], 1):
    A(f"{i}) {t}")
# ── 143번: 숫자 뒤 단위 (가짜 문제 방지) ────────────────────────
A("143. 다음 중 옳은 것은?")
A("8개월간 추적 관찰한 결과를 정리한 것이다.")
A("1. 하나")
A("2. 둘")
A("3. 셋")
A("4. 넷")
A("5. 다섯")
# ── 144번: 구분점 없음 — 사용자가 보고한 143/144 병합 지점 ───────
A("144다음 중 HIV의 진단 및 치료와 관련한 설명으로 옳지 않은 것은?")
A("1. 상용화된 백신은 아직 없다.")
A("2. 항체 검사로 진단한다.")
A("3. 역전사효소 억제제를 쓴다.")
A("4. CD4 수를 추적한다.")
A("5. 완치가 가능하다.")

d.save(OUT / "real2022.docx")


# ── 2023 기말: 번호가 띄엄띄엄하고 선지가 16개인 문제가 있다 ───────────
d = docx.Document()
A = d.add_paragraph
for t in ["2023학년도 1학기 감면 기말 시험 문제 복원-인하대학교 의학과 37기", "주의사항",
          "*문제는 번호, 마침표, 스페이스바 하고 써주세요",
          "*사진자료는 강의록에서 찾아 첨부해주세요.(필요시 편집 후 업로드 바람)",
          "______________________________________________"]:
    A(t)
# 1번 — 선지에 아무 표기가 없다
A("1. 세균배양용 검체 보관시 냉장보관하지 않고 실온에 보관하여야 하는 검체는 무엇인가?")
for t in ["농양", "뇌척수액", "소변", "객담", "인후검체"]:
    A(t)
# 22번 — 선지에 "1)" 표기
A("22. 설사변에서 적혈구를 탐식한 원충의 영양형이 관찰되었다. 이 증상의 기전은?")
for t in ["1) 장벽의 림프절을 침범하여 육아종을 일으킨다.", "2) 담즙염을 분해시켜 지방 소화를 억제한다.",
          "3) 장벽의 근육층을 약화시켜 장의 연동운동을 억제한다.",
          "4) 충체가 장 표면에 부착하여 삼투성 설사를 유발한다.",
          "5) 장 조직을 침범해 플라스크 모양의 궤양을 만든다."]:
    A(t)
# 32번 — 마침표 뒤 공백 없음
A("32.mucocutaneous leishmaniasis를 유발하는 매개체가 아닌것은?")
for t in ["1.브라질리슈만편모충", "2.멕시코리슈만편모충", "3.베네수엘라리슈만편모충",
          "4.아마존리슈만편모충", "5.소아리슈만편모충"]:
    A(t)
# 66번 — "66번." 표기 + 번호 없는 선지 6개
A("66번. 바이러스에 대한 설명으로 맞는 것을 한가지 이상 고르시오.")
for t in ["감염성 바이러스 입자를 viroid라고 함.", "음성가닥의 rna 바이러스 유전자는 감염성이 없음",
          "바이러스는 일반적으로 산성에 안정적임", "일반적으로 동물감염 바이러스는 복합구조를 가짐",
          "신종 바이러스 출현에 대비한 항생제 개발이 절실함",
          "대부분의 바이러스는 한 종의 숙주세포 수용체가 있음"]:
    A(t)
# 87번 — 선지 16개. expect<=5 규칙이 여기서 문제를 11개로 쪼개고 있었다
A("87. 종양 유발과 가장 연관된 바이러스는?")
for i, t in enumerate(["인플루엔자 1", "인플루엔자 2", "인플루엔자 3", "간염바이러스 A",
                       "간염바이러스 B", "간염바이러스 C", "간염바이러스 D", "간염바이러스 E",
                       "HPV 1,2,3,4", "HPV 6,11", "HPV 16,18", "아데노 1,2,5",
                       "아데노 3,7", "아데노 8,19,37", "아데노 40,41", "아데노 11,21"], 1):
    A(f"{i}) {t}")
# 110번 — <보기> 상자 + 조합형 선지
A("110. 벼룩이 매개하는 질병으로 올바른 것을 모두 고른 것은?")
A("<보기> a. 수면병 b. 옴 c. 개조충 d. 발진열 e. 페스트")
for t in ["1) a, b, e", "2) b, c, e", "3) a, d, e", "4) b, c, d", "5) c, d, e"]:
    A(t)
d.add_paragraph()._p.append(parse_xml(TEXTBOX))
d.save(OUT / "real_2023.docx")


# ── 합성 픽스처: 텍스트 상자 중복·보기 상자 처리 ──────────────────────
d = docx.Document()
A = d.add_paragraph
A("2023학년도 감면 기말고사")
A("1. 다음 중 옳은 것은?")
for t in ["1. 보기 하나", "2. 보기 둘", "3. 보기 셋"]:
    A(t)
A("2. 두 번째 문제이다.")
for t in ["(그림 설명: 바이러스 구조도)", "<보기>", "가. 첫째", "나. 둘째", "모두 고른 것은?"]:
    A(t)
d.add_paragraph()._p.append(parse_xml(TEXTBOX))
d.save(OUT / "t.docx")


# ── 2025 소화기: 선지에 소수, 뒤바뀐 번호, 번호만 있는 줄 ──────────────
d = docx.Document()
A = d.add_paragraph
def choices(*xs):
    for t in xs:
        p = A(t)
        p._p.append(parse_xml(NUMPR))          # 워드 자동 번호 목록 서식

A("2025 소화기 문제복원")
# 12번 — 첫 선지가 "38.5도의 발열…". 38번 문제로 잡혀 13~37번이 통째로 날아갔다
A("12. 87세 여성 환자가 3일전부터 발생한 발열과 복통을 주소로 내원했다. 옳지 않은 것은?")
choices("38.5도의 발열이 나타나므로 원인균을 찾기 위한 대변 검사를 해야한다.",
        "고령의 환자이므로 설사로 인한 탈수를 고려해 수액 투여를 해야 한다.",
        "수양성 설사이므로 경험적 항생제 투여를 시작해야 한다.",
        "수액 및 전해질 공급이 우선되어야 한다.",
        "세균 배양을 해도 원인균이 동정되지 않을 수 있다.")
A("13. 흡수 장애 환자의 진단 방법에 대한 설명으로 틀린 것은?")
choices("가", "나", "다", "라", "마")
# 14 : 소수처럼 보이지만 진짜 문제 머리 (바로 다음 번호이므로 인정한다)
A("14.3세 남아가 복통으로 내원하였다. 가장 알맞은 진단은?")
choices("가", "나", "다")
A("26. Achalasia의 병태생리로 알맞은 것은?")
choices("가", "나", "다", "라", "마")
A("27. Sliding hernia에서 나타날 수 있는 조직학적 소견은?")
choices("가", "나", "다", "라")
# 87 → 89 → 88 : 복원하다 번호 순서가 뒤바뀌었다
A("87. 면역관문억제제 사용의 좋은 소견이 아닌 것은?")
choices("MSI-h", "Tumor mutation burden", "MMR proficient")
A("89. 혀의 dorsal part이다. 다음 설명 중 옳은 것은?")
choices("(가) - keratinized stratified squamous epithelieum",
        "(나) - gustatory cell 의 microvilli 에 receptor 존재")
A("88.유전자 변이와 그에 따른 표적 치료제가 짝지어진 것중 옳지 않은 것은?")
choices("Her2 변이 - trastuzumab", "EGFR 변이 - osimertinib")
# 126 : 번호만 찍고 지문은 다음 줄에
A("126.")
A("50세 여자가 최근 피로감을 느껴 간검사를 위해 병원에 왔다. 옳은 것은?")
choices("Lamivudine", "Adefovir Dipivoxil", "Entecavir")
A("136. 세균 간 농양에 대한 설명으로 옳은 것은?")
choices("경피적 배액이 기본이다.", "K. pneumoniae 에 의해 가장 많이 발생한다.")
# 137 : 마침표 없이 번호만 붙었다
A("137 소화성 궤양의 특징으로 옳은것은?")
A("H pylori로 인한 소화성궤양은 십이지장이 위보다 많다")
d.save(OUT / "real2025.docx")

for f in sorted(os.listdir(OUT)):
    print(f"{os.path.getsize(OUT / f) / 1024:7.0f} KB  fixtures/{f}")
