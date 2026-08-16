# -*- coding: utf-8 -*-
"""
족보 문제 추출 핵심 모듈 (엑셀 불필요 버전)
- 선택한 문제만 추출하여 하나의 PDF로 병합
- 풀이 PDF가 있으면 우선 사용, 없으면 시험 DOCX에서 재렌더링
- 풀이 있는 문제: 풀이 PDF에서 (문제+해설) 페이지 블록 추출
- 풀이 없는 문제: 시험 DOCX를 PDF로 변환 후 해당 문제 페이지 추출
- 연도 최신 -> 오래된 순 정렬
"""
import os
import re
import glob
from io import BytesIO

from pypdf import PdfReader, PdfWriter


# ---------------------------------------------------------------------------
# 폰트: 깔끔한 산세리프(고딕) 우선. 시스템에 설치된 TTF를 찾고, 없으면 내장 고딕 사용.
# ---------------------------------------------------------------------------
_FONTS = None  # (regular, bold)

# (패밀리이름, 일반 TTF 경로, 볼드 TTF 경로) — 위에서부터 먼저 발견되는 것을 사용
_HERE = os.path.dirname(os.path.abspath(__file__))

_FONT_CANDIDATES = [
    # 0) 저장소에 함께 넣어둔 폰트 (가장 확실 — 어느 서버에서든 동일한 결과)
    ("KoreanUI", os.path.join(_HERE, "fonts", "NanumBarunGothic.ttf"),
     os.path.join(_HERE, "fonts", "NanumBarunGothicBold.ttf")),
    ("KoreanUI", os.path.join(_HERE, "fonts", "NanumGothic.ttf"),
     os.path.join(_HERE, "fonts", "NanumGothicBold.ttf")),
    # 1) Linux - 나눔바른고딕 (본문 가독성이 좋아 최우선)
    ("KoreanUI", "/usr/share/fonts/truetype/nanum/NanumBarunGothic.ttf",
     "/usr/share/fonts/truetype/nanum/NanumBarunGothicBold.ttf"),
    ("KoreanUI", "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
     "/usr/share/fonts/truetype/nanum/NanumGothicBold.ttf"),
    # 2) Windows - 맑은 고딕
    ("KoreanUI", r"C:\Windows\Fonts\malgun.ttf", r"C:\Windows\Fonts\malgunbd.ttf"),
    ("KoreanUI", r"C:\Windows\Fonts\NanumBarunGothic.ttf",
     r"C:\Windows\Fonts\NanumBarunGothicBold.ttf"),
    ("KoreanUI", r"C:\Windows\Fonts\NanumGothic.ttf", r"C:\Windows\Fonts\NanumGothicBold.ttf"),
    # 3) macOS - Apple SD 산돌고딕 Neo
    ("KoreanUI", "/System/Library/Fonts/AppleSDGothicNeo.ttc", None),
    ("KoreanUI", "/Library/Fonts/AppleSDGothicNeo.ttc", None),
    # 4) Noto Sans CJK
    ("KoreanUI", "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
     "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc"),
    ("KoreanUI", "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
     "/usr/share/fonts/truetype/noto/NotoSansCJK-Bold.ttc"),
]


def _ensure_fonts():
    """(일반, 볼드) 폰트 이름을 반환. 깔끔한 고딕 계열을 우선 사용."""
    global _FONTS
    if _FONTS is not None:
        return _FONTS
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont

    for name, reg, bold in _FONT_CANDIDATES:
        try:
            if not reg or not os.path.exists(reg):
                continue
            sub = 0 if reg.lower().endswith(".ttc") else None
            pdfmetrics.registerFont(TTFont(name, reg, subfontIndex=sub) if sub is not None
                                    else TTFont(name, reg))
            bold_name = name
            if bold and os.path.exists(bold):
                bsub = 0 if bold.lower().endswith(".ttc") else None
                pdfmetrics.registerFont(TTFont(name + "-Bold", bold, subfontIndex=bsub)
                                        if bsub is not None else TTFont(name + "-Bold", bold))
                bold_name = name + "-Bold"
            pdfmetrics.registerFontFamily(
                name, normal=name, bold=bold_name, italic=name, boldItalic=bold_name)
            _FONTS = (name, bold_name)
            return _FONTS
        except Exception:
            continue

    # 폴백: reportlab 내장 한국어 고딕. 별도 파일이 필요 없지만
    # 영문이 Helvetica로 따로 렌더링되어 한글과 섞이면 모양이 떨어진다.
    # (되도록 fonts/ 폴더에 NanumBarunGothic.ttf 를 넣어 이 경로를 피할 것)
    try:
        pdfmetrics.registerFont(UnicodeCIDFont("HYGothic-Medium"))
        pdfmetrics.registerFontFamily(
            "HYGothic-Medium", normal="HYGothic-Medium", bold="HYGothic-Medium",
            italic="HYGothic-Medium", boldItalic="HYGothic-Medium")
        _FONTS = ("HYGothic-Medium", "HYGothic-Medium")
    except Exception:
        _FONTS = ("Helvetica", "Helvetica-Bold")
    return _FONTS


def font_status():
    """현재 사용 중인 폰트 정보를 (이름, 실제_TTF경로_또는_None) 로 반환. UI 안내용."""
    for name, reg, _bold in _FONT_CANDIDATES:
        if reg and os.path.exists(reg):
            return name, reg
    return "HYGothic-Medium(내장)", None


def _ensure_title_font():
    return _ensure_fonts()[0]


def make_title_page(title, subtitle=""):
    """연도 구분 표지. 가로로 긴(레터박스) 비율의 깔끔한 배너 1장을 반환."""
    from reportlab.pdfgen import canvas
    from reportlab.lib.colors import HexColor

    font, font_bold = _ensure_fonts()

    # 가로로 긴 비율 (약 2.4 : 1)
    w, h = 900.0, 380.0
    # 색상 팔레트
    bg = HexColor("#F4F7FB")       # 아주 옅은 배경
    card = HexColor("#FFFFFF")     # 카드
    accent = HexColor("#2E6CA4")   # 포인트 남색
    accent2 = HexColor("#8FB4D6")  # 옅은 남색
    ink = HexColor("#1F2933")      # 제목 글자
    muted = HexColor("#7B8794")    # 부제 글자

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=(w, h))

    # 배경
    c.setFillColor(bg)
    c.rect(0, 0, w, h, stroke=0, fill=1)

    # 가운데 카드 (은은한 테두리)
    m = 34
    c.setFillColor(card)
    c.setStrokeColor(HexColor("#E1E7EF"))
    c.setLineWidth(1)
    c.roundRect(m, m, w - 2 * m, h - 2 * m, 18, stroke=1, fill=1)

    # 왼쪽 세로 포인트 바
    barx = m + 40
    c.setFillColor(accent)
    c.roundRect(barx, m + 46, 8, h - 2 * m - 92, 4, stroke=0, fill=1)

    # 텍스트 왼쪽 정렬 기준선
    tx = barx + 34
    mid = h / 2

    # 상단 작은 라벨 (부제)
    if subtitle:
        c.setFillColor(accent)
        c.setFont(font, 16)
        c.drawString(tx, mid + 66, subtitle)

    # 연도 (크게, 볼드)
    c.setFillColor(ink)
    c.setFont(font_bold, 96)
    c.drawString(tx - 4, mid - 34, str(title))

    # 연도 아래 포인트 밑줄 (두 톤)
    uy = mid - 58
    c.setFillColor(accent)
    c.roundRect(tx, uy, 118, 5, 2.5, stroke=0, fill=1)
    c.setFillColor(accent2)
    c.roundRect(tx + 124, uy, 40, 5, 2.5, stroke=0, fill=1)

    # 오른쪽 아래 작은 문구
    c.setFillColor(muted)
    c.setFont(font, 12)
    c.drawRightString(w - m - 40, m + 34, "기출 문제 발췌")

    c.showPage()
    c.save()
    buf.seek(0)
    return PdfReader(buf).pages[0]


# ---------------------------------------------------------------------------
# 파일 이름 규칙
#   시험 문제 :  "{연도} {과목} {중간기말}.docx"          (예: 2023 감면 기말.docx)
#   풀이       :  "{연도} {과목} {중간기말} 풀이.pdf"      (예: 2023 감면 기말 풀이.pdf)
# ---------------------------------------------------------------------------


def scan_library(folder):
    """자료 폴더에 등록된 파일을 훑어 시험 세트 목록을 반환.

    반환: [{"year": int, "term": str, "subject": str,
            "solution": 파일명 or None, "docx": 파일명 or None}, ...]
    (연도 최신 -> 오래된 순)
    """
    sets = {}   # (year, term) -> dict
    if not os.path.isdir(folder):
        return []

    for path in sorted(glob.glob(os.path.join(folder, "*"))):
        name = os.path.basename(path)
        if name.startswith("~$") or name.startswith("."):
            continue
        ext = os.path.splitext(name)[1].lower()
        if ext not in (".pdf", ".docx"):
            continue

        m = re.search(r"(19|20)\d{2}", name)
        if not m:
            continue
        year = int(m.group(0))
        term = "중간" if "중간" in name else ("기말" if "기말" in name else None)
        if term is None:
            continue

        # 과목명 추정: "{연도} {과목} {학기}" 형태에서 가운데 토큰
        subject = ""
        mm = re.search(r"(?:19|20)\d{2}\s+(.+?)\s+(?:중간|기말)", name)
        if mm:
            subject = mm.group(1).strip()

        # 과목까지 키에 넣는다. (연도, 학기)만으로 묶으면 같은 연도·학기의 다른
        # 과목 파일이 서로 덮어써서, 표시된 과목명과 실제 파일이 어긋난다.
        key = (year, term, subject)
        rec = sets.setdefault(key, {
            "year": year, "term": term, "subject": subject,
            "solution": None, "docx": None,
        })

        if ext == ".pdf" and "풀이" in name:
            rec["solution"] = name
        elif ext == ".docx" and "풀이" not in name:
            rec["docx"] = name

    return sorted(sets.values(),
                  key=lambda r: (-r["year"], _TERM_ORDER.get(r["term"], 9),
                                 r["subject"]))


def find_source(folder, year, subject, term, kind):
    """kind='풀이'(pdf) 또는 '문제'(docx) 원본 파일 경로를 찾는다."""
    if kind == "풀이":
        exact = os.path.join(folder, f"{year} {subject} {term} 풀이.pdf")
        if os.path.exists(exact):
            return exact
        for p in glob.glob(os.path.join(folder, "*.pdf")):
            b = os.path.basename(p)
            # 과목까지 확인한다. 빠뜨리면 요청한 과목이 없을 때 다른 과목 파일을
            # 조용히 반환해 엉뚱한 문제가 결과에 섞인다.
            if str(year) in b and term in b and "풀이" in b and subject in b:
                return p
        return None
    else:  # 문제 docx
        for ext in (".docx", ".DOCX"):
            exact = os.path.join(folder, f"{year} {subject} {term}{ext}")
            if os.path.exists(exact):
                return exact
        for p in glob.glob(os.path.join(folder, "*.docx")):
            b = os.path.basename(p)
            if b.startswith("~$"):
                continue
            if str(year) in b and term in b and "풀이" not in b and subject in b:
                return p
        return None


# ---------------------------------------------------------------------------
# 풀이 PDF: 문제별 페이지 블록 매핑
# ---------------------------------------------------------------------------
# 풀이 슬라이드의 "(1/3)" 같은 진행 표기 (있으면 쓰고, 없어도 동작한다)
_SLIDE_MARK_RE = re.compile(r"\(\s*(\d{1,2})\s*/\s*(\d{1,2})\s*\)")
# 슬라이드 상단의 문제 번호. 글머리 기호가 앞에 붙거나 마침표가 없는 경우도 잡는다.
#   숫자는 '온전한 수'여야 한다. (?<!\d)...(?!\d) 가 없으면 표지의 "2023" 에서
#   앞 세 자리 "202" 를 문제 번호로 잘라 읽고, last_q 가 202 로 올라가
#   이후의 진짜 문제 번호가 전부 무시된다 (풀이 PDF 전체가 한 블록이 됨).
#   "48. influenza virus..." / "75 그림과관련된질병..." / "• 76 사람면역결핍..."
#   구분자('.' 또는 ')')를 요구한다. 없으면 표지의 "35기 족보"가 35번으로
#   잡히고, last_q 가 35 로 올라가 1~34번이 전부 사라진다.
_SOL_QNUM_RE = re.compile(
    r"(?:^|\n)\s*(?:[•·\-\*\u2022]\s*)?(?<!\d)(\d{1,3})(?!\d)\s*[.)]\s*(\S[^\n]{3,})")
#   구분자를 아예 쓰지 않는 족보("75 그림과관련된질병")를 위한 완화판.
#   strict 패턴으로 거의 못 찾았을 때만 대체로 쓴다.
_SOL_QNUM_RE_LOOSE = re.compile(
    r"(?:^|\n)\s*(?:[•·\-\*\u2022]\s*)?(?<!\d)(\d{1,3})(?!\d)\s*[.)]?\s*(\S[^\n]{3,})")
# "정답: 2번" 처럼 해설 슬라이드임을 알려주는 머리글
_ANSWER_HEAD_RE = re.compile(r"^\s*(정답|답)\s*[:：）)]?")
# 문제 번호는 슬라이드 상단에 나온다. 아래쪽의 "77. 인체감염진균"(강의록 참조)이
# 문제 번호로 오인되지 않도록 탐색 범위를 앞부분 몇 줄로 제한한다.
_SOL_HEAD_LINES = 4


def _solution_page_texts(pdf_path):
    """풀이 PDF의 페이지별 텍스트. PyMuPDF 가 있으면 그쪽을 쓴다.

    pypdf 의 추출은 줄바꿈이 뭉개져 문제 번호를 놓치는 경우가 많아
    (같은 파일에서 78문제 -> 64문제로 감소) PyMuPDF 를 우선 사용한다.
    """
    try:
        try:
            import pymupdf as fitz   # PyMuPDF (신형 이름)
        except ImportError:
            import fitz              # 구버전 호환
        with fitz.open(pdf_path) as doc:
            return [pg.get_text() or "" for pg in doc]
    except Exception:
        reader = PdfReader(pdf_path)
        return [(pg.extract_text() or "") for pg in reader.pages]


def _scan_solution_starts(texts, qnum_re):
    """페이지 텍스트 목록에서 [(페이지 index, 문제번호), ...] 를 찾는다."""
    starts = []
    last_q = 0
    for i, t in enumerate(texts):
        mk = _SLIDE_MARK_RE.search(t)
        k = int(mk.group(1)) if mk else None
        if k is not None and k > 1:
            continue                      # 이어지는 슬라이드
        if _ANSWER_HEAD_RE.match(t.strip()):
            continue                      # 해설 슬라이드
        head_lines = [ln for ln in t.splitlines() if ln.strip()][:_SOL_HEAD_LINES]
        head = "\n" + "\n".join(head_lines)
        pick = None
        for m in qnum_re.finditer(head):
            num = int(m.group(1))
            if last_q < num <= 400:
                pick = num
                break
        if pick is None:
            continue
        starts.append((i, pick))
        last_q = pick
    return starts


def map_solution_blocks(pdf_path):
    """풀이 PDF에서 {문제번호: (start_page, end_page_exclusive)} 반환.

    한 문제는 보통 '문제 / 정답·해설 / 강의록' 슬라이드로 이어진다.
    족보 파일이 규칙을 잘 지키지 않으므로 여러 단서를 함께 본다.

      - (1/3) 표기가 있으면 활용하되, **없어도 동작한다.**
        (표기가 전혀 없는 족보에서 문제가 하나도 안 잡히는 문제가 있었다)
      - 표기가 (2/3) 처럼 첫 장이 아니면 이어지는 슬라이드로 본다.
      - '정답'으로 시작하는 슬라이드는 해설이므로 블록 시작이 아니다.
      - 문제 번호는 '앞 문제보다 큰 번호'만 인정한다.
        이렇게 하면 선지 번호(1~5)나 반복 표시된 같은 번호에 속지 않는다.
    """
    reader = PdfReader(pdf_path)
    texts = _solution_page_texts(pdf_path)
    n = len(texts)

    starts = _scan_solution_starts(texts, _SOL_QNUM_RE)
    if len(starts) < 2:
        # 구분자 없는 족보일 수 있다. 완화판으로 다시 훑어 더 많이 잡히면 그쪽을 쓴다.
        loose = _scan_solution_starts(texts, _SOL_QNUM_RE_LOOSE)
        if len(loose) > len(starts):
            starts = loose

    blocks = {}
    for idx, (pg, num) in enumerate(starts):
        end = starts[idx + 1][0] if idx + 1 < len(starts) else n
        blocks.setdefault(num, (pg, end))
    return blocks, reader


# ---------------------------------------------------------------------------
# 시험 DOCX 파싱 (풀이 없는 문제) -> 문제별로 직접 파싱 후 깔끔하게 재렌더링
# ---------------------------------------------------------------------------
# 숫자 표기 줄:  "12." / "3)" / "91번."  ->  (번호, 나머지 텍스트)
_NUMBERED_RE = re.compile(r"^\s*(\d{1,3})\s*번?\s*[.)．]\s*(.*)$")
_DASH_RE = re.compile(r"^\s*[-–—•*·]\s*(.+)$")
_CIRC_RE = re.compile(r"^\s*([①②③④⑤⑥⑦⑧⑨⑩])\s*(.+)$")
# <보기> 항목 표기:  가. 나. ... / ㄱ. ㄴ. ...  (모두 고른 것 유형의 '보기')
_KOX_RE = re.compile(r"^\s*[가나다라마바사아자차]\s*[.)]\s*(.+)$")
_JAMO_RE = re.compile(r"^\s*[ㄱ-ㅎ]\s*[.)]\s*(.+)$")
_MAX_CHOICES = 5   # 선지는 최대 5개(1~5)로 가정 -> 문제번호와의 혼동 방지


def _is_bogi_header(line):
    """'<보기>', '(보기 ...)', '[보기]', '보기)' 등 보기 상자 시작 줄인지."""
    s = line.replace(" ", "")
    return (s.startswith("<보기") or s.startswith("(보기") or s.startswith("[보기")
            or s.startswith("보기>") or s.startswith("보기]") or s.startswith("<보 기"))


def _is_annotation(line):
    """줄 전체가 괄호로 묶인 '정리자 주석'인지.

    족보 파일에는 원본 문제에 없던 설명이 괄호로 덧붙는 경우가 많다.
      예) "(H3N2와 H1N1 바이러스가 유전자 재조합하는 그림)"
          "(보기 8개 중, 정답 4개 고르는 문제)"
    이런 줄은 선지가 아니라 안내문이므로 번호를 붙이면 안 된다.
    """
    s = line.strip()
    return (len(s) >= 2
            and s[0] in "(（[" and s[-1] in ")）]"
            and not _is_bogi_item(s))


def _is_bogi_item(line):
    """'가.' '나.' … / 'ㄱ.' 'ㄴ.' … 처럼 '보기 항목'(단일 라벨)인지."""
    return bool(_KOX_RE.match(line) or _JAMO_RE.match(line))


def _strip_marker(line):
    """줄 앞의 선지/보기 기호(-, 1., 3), ①, 가., ㄱ.)를 떼고 본문만 반환."""
    for rx in (_DASH_RE, _CIRC_RE, _NUMBERED_RE, _KOX_RE, _JAMO_RE):
        m = rx.match(line)
        if m:
            return m.group(m.lastindex).strip()
    return line


def _classify_body(lines):
    """지문 이후 줄들을 (제시문 줄, 실제 선지) 로 분리한다.

    방침: '선지'를 양성적으로 인식해서 그 줄만 번호를 매긴다.
      나머지(보기·안내문·지문 연장 등 제시문)는 형태를 추측하지 않고 그대로 둔다.

    1) 선지 표기가 있으면(대시 '-', 동그라미 '①', 숫자 '1.'/'1)') 그 표기의 줄만 선지.
       -> 표기 없는 줄(보기 상자, '하나를 선택하세요' 등)은 모두 제시문.
    2) 어떤 선지 표기도 없으면('모두 고른 것' 조합형):
       -> 보기 항목(가./ㄱ.)·보기 상자는 제시문, 나머지 줄이 선지.
       -> 단, 보기 항목만 있으면(가/나/다 단일선택) 그 자체가 선지.

    반환: (presented, choices)
      presented = [(줄 텍스트, is_bogi_bool), ...]   # 번호 없이 그대로 표시
      choices   = [선지 텍스트, ...]                 # 렌더 시 1) 2) 3) ...
    """
    # 입력은 (줄, 워드_목록서식_여부) 튜플. 예전 형식(문자열)도 받아들인다.
    items = []
    for it in lines:
        if isinstance(it, tuple):
            txt, is_list = it
        else:
            txt, is_list = it, False
        if txt:
            items.append((txt, is_list))
    if not items:
        return [], []
    lines = [t for t, _ in items]

    # 선지에 쓰인 표기를 감지 (대시 > 동그라미 > 숫자 순)
    if any(_DASH_RE.match(l) for l in lines):
        marker_rx = _DASH_RE
    elif any(_CIRC_RE.match(l) for l in lines):
        marker_rx = _CIRC_RE
    elif any(_NUMBERED_RE.match(l) for l in lines):
        marker_rx = _NUMBERED_RE
    else:
        marker_rx = None

    presented, choices = [], []

    if marker_rx is not None:
        for l in lines:
            m = marker_rx.match(l)
            if m:
                choices.append(m.group(m.lastindex).strip())   # 표기 있는 줄 = 선지
            else:
                presented.append((l, False))   # 안내문·보기 모두 그대로 표시
        return presented, choices

    # 글자 표기는 없지만 워드에서 '자동 번호 목록'으로 서식이 지정된 줄이 있으면
    # 그 줄들만 선지로 본다. (그림 설명·안내문 같은 일반 문단이 선지로 섞이는 것을 방지)
    if any(is_list for _t, is_list in items):
        for txt, is_list in items:
            if is_list:
                choices.append(_strip_marker(txt))
            else:
                presented.append((txt, False))   # 안내문·보기 모두 그대로 표시
        return presented, choices

    # 표기가 전혀 없는 경우 ('모두 고른 것' 조합형 등)
    #  - 괄호 안내문은 선지 후보에서 제외한다.
    #  - 이미 자기 라벨(가./ㄱ.)을 가진 보기 항목은 선지로 번호를 다시 매기지 않고
    #    원래 라벨 그대로 보여준다. (선지 개수로 추측하지 않으므로 선지가 5개를
    #    넘는 문제에서도 잘못 분류되지 않는다)
    for l in lines:
        if _is_annotation(l):
            presented.append((l, False))
        elif _is_bogi_item(l) or _is_bogi_header(l):
            presented.append((l, False))
        else:
            choices.append(_strip_marker(l))
    return presented, choices


def _iter_body_blocks(doc):
    """본문을 '문단' 또는 '표' 단위로 문서에 보이는 순서대로 반환.

    ("p", Paragraph) 또는 ("tbl", Table) 튜플을 내놓는다.

    python-docx 의 doc.paragraphs 는 최상위 문단만 주므로 표·텍스트 상자 안의
    내용이 누락된다. 반대로 모든 w:p 를 평평하게 훑으면 표 구조가 사라져
    나란히 배치돼야 할 칸이 위아래로 쌓인다. 그래서 표는 표로 따로 넘긴다.
    """
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    body = doc.element.body
    p_tag = qn("w:p")
    tbl_tag = qn("w:tbl")
    # mc(markup compatibility) 는 python-docx 기본 네임스페이스 맵에 없어 직접 지정
    fallback_tag = "{http://schemas.openxmlformats.org/markup-compatibility/2006}Fallback"

    def _ancestor_tags(el):
        """el 의 조상 태그들을 body 까지 거슬러 올라가며 내놓는다.

        lxml 은 같은 XML 요소에 대해 매번 새 파이썬 객체를 만들 수 있어
        id() 비교가 어긋난다. 그래서 부모 연결을 직접 확인한다.
        """
        cur_el = el.getparent()
        while cur_el is not None and cur_el is not body:
            yield cur_el.tag
            cur_el = cur_el.getparent()

    for el in body.iter(p_tag, tbl_tag):
        anc = set(_ancestor_tags(el))
        # mc:Fallback 하위 내용은 mc:Choice 와 중복되므로 제외
        if fallback_tag in anc:
            continue
        # 표 안의 문단·중첩 표는 표 블록에서 함께 처리하므로 개별로 내보내지 않는다
        if tbl_tag in anc:
            continue
        if el.tag == tbl_tag:
            tbl = Table(el, doc)
            # 칸이 하나뿐인 표는 사실상 '테두리 있는 문단'이므로 일반 흐름으로 처리
            try:
                single = len(tbl.rows) == 1 and len(tbl.rows[0].cells) == 1
            except Exception:
                single = False
            if single:
                for cp in tbl.rows[0].cells[0].paragraphs:
                    yield "p", cp
            else:
                yield "tbl", tbl
        else:
            yield "p", Paragraph(el, doc)


def _table_spec(tbl):
    """표를 렌더링용 자료구조로 변환.

    반환: [[{"text": 셀텍스트, "images": [blob, ...]}, ...], ...]  (행 x 열)
    """
    rows = []
    for r in tbl.rows:
        cells = []
        seen = set()
        for c in r.cells:
            # 병합된 칸은 같은 객체가 반복 등장하므로 한 번만 담는다
            key = id(c._tc)
            if key in seen:
                continue
            seen.add(key)
            texts, imgs = [], []
            for p in c.paragraphs:
                t = p.text.strip()
                if t:
                    texts.append(t)
                imgs.extend(_para_images(p))
            cells.append({"text": "\n".join(texts), "images": imgs})
        if cells:
            rows.append(cells)
    return rows


def _para_images(para):
    """문단(paragraph)에 포함된 이미지 blob 리스트를 반환.

    DrawingML(a:blip) 뿐 아니라 구형 VML(v:imagedata) 이미지도 함께 처리한다.
    """
    from docx.oxml.ns import qn
    imgs = []
    seen = set()

    def _add(rid):
        if not rid or rid in seen:
            return
        seen.add(rid)
        try:
            part = para.part.related_parts[rid]
            imgs.append(part.blob)
        except Exception:
            pass

    # 최신 형식 (DrawingML) — inline / anchor(떠 있는 그림) 모두 포함
    for blip in para._element.findall(".//" + qn("a:blip")):
        _add(blip.get(qn("r:embed")) or blip.get(qn("r:link")))

    # 구형 형식 (VML) — 예전 워드에서 삽입된 그림
    # v(VML) 네임스페이스는 python-docx 기본 맵에 없어 직접 지정한다
    vml_tag = "{urn:schemas-microsoft-com:vml}imagedata"
    for vml in para._element.findall(".//" + vml_tag):
        _add(vml.get(qn("r:id")) or vml.get(qn("r:href")))

    return imgs


def _is_list_paragraph(para):
    """워드의 '자동 번호/글머리 기호 목록'으로 서식이 지정된 문단인지 판별.

    선지를 자동 번호 목록으로 만든 경우 번호가 텍스트에 없어서,
    그림 설명 같은 일반 문단과 구분되지 않는다. 서식(w:numPr)을 직접 확인한다.
    """
    from docx.oxml.ns import qn
    ppr = para._element.find(qn("w:pPr"))
    if ppr is not None and ppr.find(qn("w:numPr")) is not None:
        return True
    # 문단에 직접 지정되지 않고 스타일에 정의된 경우(예: List Number 스타일)
    try:
        style = para.style
        while style is not None:
            spr = style.element.find(qn("w:pPr"))
            if spr is not None and spr.find(qn("w:numPr")) is not None:
                return True
            style = style.base_style
    except Exception:
        pass
    return False


def _para_lines(para):
    """문단을 '물리적 줄' 단위로 분해. 한 문단 안의 줄바꿈(<w:br/>)·탭도 반영한다.
    (python-docx의 paragraph.text 는 문단 내 줄바꿈을 무시하므로 직접 처리)"""
    from docx.oxml.ns import qn
    t_tag, br_tag, cr_tag, tab_tag = (
        qn("w:t"), qn("w:br"), qn("w:cr"), qn("w:tab"))
    # mc:Fallback 은 mc:Choice 와 같은 내용을 중복해 담고, w:txbxContent(텍스트
    # 상자)는 _iter_body_blocks 가 별도 문단으로 따로 내보낸다. 두 갈래로 모두
    # 내려가면 같은 문장이 2~3번 출력된다.
    fallback_tag = "{http://schemas.openxmlformats.org/markup-compatibility/2006}Fallback"
    txbx_tag = qn("w:txbxContent")
    parts = []

    def walk(el):
        for node in el:
            tag = node.tag
            if tag == fallback_tag or tag == txbx_tag:
                continue
            if tag == t_tag:
                parts.append(node.text or "")
            elif tag in (br_tag, cr_tag):
                parts.append("\n")
            elif tag == tab_tag:
                parts.append(" ")
            walk(node)

    walk(para._element)
    text = "".join(parts)
    return [ln.strip() for ln in text.split("\n")]


def compose_slides_grid(src_pages, label, per_page=4):
    """풀이 슬라이드 여러 장을 한 페이지에 모아 배치한 페이지들을 반환.

    한 문제(문제+해설)를 한 페이지에서 보기 위한 기능이다.
    - A4 가로 방향에 2x2로 최대 4장 배치
    - 슬라이드가 4장을 넘으면 4장씩 나눠 여러 페이지로 (5~8장 -> 2페이지)
    - 왼쪽 위에 '2023 기말 48번' 같은 라벨을 얹는다
    """
    from pypdf import PageObject, Transformation

    W, H = 841.89, 595.28          # A4 가로
    mx, my = 16, 14                # 좌우 / 위아래 여백
    top = 16                       # 라벨 자리
    gap = 8
    cols, rows = 2, 2
    cw = (W - 2 * mx - (cols - 1) * gap) / cols
    ch = (H - 2 * my - top - (rows - 1) * gap) / rows

    out = []
    for chunk_i in range(0, len(src_pages), per_page):
        chunk = src_pages[chunk_i:chunk_i + per_page]
        page = PageObject.create_blank_page(width=W, height=H)
        for i, src in enumerate(chunk):
            try:
                box = src.mediabox
                x0, y0 = float(box.left), float(box.bottom)
                sw = float(box.width) or 1.0
                sh = float(box.height) or 1.0
                s = min(cw / sw, ch / sh)
                col, row = i % cols, i // cols
                # PDF 좌표계는 왼쪽 아래가 원점
                cell_x = mx + col * (cw + gap)
                cell_y = H - my - top - (row + 1) * ch - row * gap
                tx = cell_x + (cw - sw * s) / 2 - x0 * s
                ty = cell_y + (ch - sh * s) / 2 - y0 * s
                page.merge_transformed_page(
                    src, Transformation().scale(s).translate(tx, ty))
            except Exception:
                continue
        if label:
            try:
                page.merge_page(_make_label_page(label, W, H, mx, my))
            except Exception:
                pass
        out.append(page)
    return out


def _make_label_page(text, W, H, mx, my):
    """페이지 왼쪽 위에 얹을 라벨(문제 번호) 페이지를 만든다."""
    from reportlab.pdfgen import canvas
    _font, font_bold = _ensure_fonts()
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=(W, H))
    c.setFillColorRGB(0.18, 0.42, 0.64)
    c.setFont(font_bold, 9)
    c.drawString(mx, H - my - 10, text)
    c.setStrokeColorRGB(0.87, 0.90, 0.94)
    c.setLineWidth(0.4)
    c.line(mx, H - my - 14, W - mx, H - my - 14)
    c.save()
    buf.seek(0)
    return PdfReader(buf).pages[0]


def parse_docx_questions(docx_path):
    """DOCX를 파싱해 {문제번호: {'num','stem','choices':[...],'images':[blob,...]}} 반환.

    핵심 난점: 선지도 '1.' '2.' 처럼 문제 번호와 똑같은 '숫자.' 형식일 수 있어,
    선지를 새 문제로 오인하기 쉽다. 그래서 상태 기계로 구분한다.
      - 지문 뒤에서 1,2,3,... 로 이어지는 '숫자.' 는 선지(최대 5개)
      - 그 외의 '숫자.' 는 새 문제 번호
      - 숫자 없는 기호(-, ①, ㄱ.)나 표시가 전혀 없는 줄도 선지로 인식
    """
    import docx
    d = docx.Document(docx_path)
    questions = {}
    cur = None
    expect = 1        # 현재 문제에서 다음에 올 '숫자.' 선지 번호(1부터)
    max_num = 0       # 지금까지 등장한 최대 문제 번호(문제 번호는 문서 내에서 증가)

    def start_question(num, stem):
        nonlocal cur, expect, max_num
        cur = {"num": num, "stem": stem.strip(), "_body": [], "images": [],
               "tables": []}
        questions[num] = cur
        expect = 1
        max_num = max(max_num, num)

    for kind, obj in _iter_body_blocks(d):
        if kind == "tbl":
            # 표는 구조를 유지해 현재 문제에 붙인다
            if cur is not None:
                spec = _table_spec(obj)
                if spec:
                    cur["tables"].append(spec)
            continue
        para = obj
        lines = _para_lines(para)
        imgs = _para_images(para)
        if not any(lines) and not imgs:
            continue
        # 워드 자동 번호 목록으로 서식이 지정된 문단인지 (선지 판별의 강력한 단서)
        is_list = _is_list_paragraph(para)
        for ln in lines:
            if not ln:
                continue
            nm = _NUMBERED_RE.match(ln)
            if nm:
                num = int(nm.group(1))
                if cur is not None and num == expect and expect <= _MAX_CHOICES:
                    # 지문 뒤에서 1,2,3.. 로 이어지는 숫자 -> 선지(본문에 원문 그대로 보관)
                    cur["_body"].append((ln, is_list))
                    expect += 1
                elif (cur is None or num > max_num) and 1 <= num <= 400:
                    # 문제 번호는 문서 내에서 커지기만 한다 -> 그럴 때만 새 문제
                    start_question(num, nm.group(2).strip())
                elif cur is not None:
                    cur["_body"].append((ln, is_list))
                continue
            # 숫자 없는 줄
            if cur is None:
                continue
            cur["_body"].append((ln, is_list))
        if cur is not None and imgs:
            cur["images"].extend(imgs)

    # 본문을 제시문(보기·안내문 등)과 실제 선지로 분리
    for q in questions.values():
        q["presented"], q["choices"] = _classify_body(q.pop("_body"))
    return questions


def _build_table_flowable(spec, avail_w, cell_style, esc):
    """표 자료구조를 reportlab Table flowable 로 변환.

    - 칸 안의 그림도 그대로 넣는다 (그림이 선지인 문제가 있음)
    - 열 너비는 사용 가능한 폭을 균등 분할
    - 글자가 전혀 없고 그림만 있는 표는 테두리를 그리지 않는다
      (그림을 나란히 놓기 위한 '배치용 표'인 경우가 많음)
    """
    from reportlab.platypus import Table as RLTable, TableStyle, Paragraph, Image
    from reportlab.lib.utils import ImageReader
    from reportlab.lib import colors

    if not spec:
        return None
    ncols = max(len(r) for r in spec)
    if ncols == 0:
        return None
    col_w = avail_w / ncols
    has_text = any(c.get("text") for r in spec for c in r)

    data = []
    for row in spec:
        out = []
        for c in row:
            items = []
            if c.get("text"):
                items.append(Paragraph(esc(c["text"]).replace("\n", "<br/>"), cell_style))
            for blob in c.get("images", []):
                try:
                    ir = ImageReader(BytesIO(blob))
                    iw, ih = ir.getSize()
                    w = min(col_w - 10, iw * 0.75)
                    h = w * ih / iw
                    items.append(Image(BytesIO(blob), width=w, height=h))
                except Exception:
                    pass
            if not items:
                items = [""]
            out.append(items if len(items) > 1 else items[0])
        # 행마다 열 수가 다를 수 있어 빈 칸으로 채운다
        out += [""] * (ncols - len(out))
        data.append(out)

    tbl = RLTable(data, colWidths=[col_w] * ncols, hAlign="LEFT")
    style = [
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]
    if has_text:
        style += [
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#C9D3DE")),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F5F9")),
        ]
    tbl.setStyle(TableStyle(style))
    return tbl


def render_docx_questions_to_pages(qdatas, header):
    """선택된 문제(qdatas: 파싱된 dict 리스트)를 가독성 있게 재렌더링한 PDF의 페이지들을 반환.
    - 선지는 1) 2) 3) 형식
    - 문제 사이에 넉넉한 줄바꿈
    - 그림(있으면) 지문 아래 삽입"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Image, HRFlowable, KeepTogether,
    )
    from reportlab.lib.utils import ImageReader

    font, font_bold = _ensure_fonts()

    # 한글 조판 공통 설정
    #  - wordWrap="CJK": 한글은 단어 단위가 아니라 글자 단위로 줄바꿈해야 자연스럽다.
    #    (미설정 시 공백 없는 긴 한글이 여백을 넘거나 줄 끝이 들쭉날쭉해진다)
    #  - leading: 글자 크기의 약 1.65배. 한글은 라틴보다 행간을 넉넉히 줘야 읽기 편하다.
    #  - alignment=0(왼쪽): 양쪽 정렬은 한글에서 낱말 사이가 벌어져 자간이 불균일해 보인다.
    CJK = {"wordWrap": "CJK", "alignment": 0, "splitLongWords": True}

    stem_style = ParagraphStyle(
        "stem", fontName=font, fontSize=11.5, leading=19, spaceAfter=6,
        textColor="#1A2028", **CJK)
    # 선지: 번호는 지문 텍스트와 같은 위치에서 시작하고,
    #       둘째 줄부터는 선지 본문에 맞춰 들여쓴다(매달린 들여쓰기).
    choice_style = ParagraphStyle(
        "choice", fontName=font, fontSize=11, leading=17.5,
        leftIndent=36, firstLineIndent=-16, spaceAfter=2,
        textColor="#2B3440", **CJK)
    note_style = ParagraphStyle(
        "note", fontName=font, fontSize=11, leading=18, spaceBefore=1, spaceAfter=4,
        textColor="#1A2028", **CJK)
    head_style = ParagraphStyle(
        "head", fontName=font_bold, fontSize=9.5, leading=13,
        textColor="#5B7A99", **CJK)

    def esc(s):
        return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm)
    avail_w = doc.width
    story = []
    if header:
        story.append(Paragraph(esc(header), head_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color="#cccccc",
                                 spaceBefore=2, spaceAfter=10))

    # 문제 지문: 번호를 왼쪽으로 내밀고 둘째 줄부터 지문에 맞춰 정렬
    stem_hang = ParagraphStyle(
        "stemHang", parent=stem_style, leftIndent=20, firstLineIndent=-20)

    for qi, qd in enumerate(qdatas):
        block = []
        if qi > 0:
            # 문제 사이 얇은 구분선 — 여백만으로 나누면 경계가 흐릿하다
            block.append(HRFlowable(width="100%", thickness=0.4, color="#E8ECF1",
                                    spaceBefore=0, spaceAfter=13))
        block.append(Paragraph(
            f"<font name='{font_bold}' color='#2E6CA4'>{qd['num']}.</font>&nbsp;"
            f"{esc(qd['stem'])}",
            stem_hang))
        for blob in qd.get("images", []):
            try:
                ir = ImageReader(BytesIO(blob))
                iw, ih = ir.getSize()
                w = min(avail_w, iw * 0.75)
                h = w * ih / iw
                max_h = 150 * mm
                if h > max_h:
                    h = max_h
                    w = h * iw / ih
                block.append(Spacer(1, 4))
                block.append(Image(BytesIO(blob), width=w, height=h))
            except Exception:
                pass
        # 표 — 원본의 행/열 구조를 유지해 그린다
        for spec in qd.get("tables", []):
            flow = _build_table_flowable(spec, avail_w, note_style, esc)
            if flow is not None:
                block.append(Spacer(1, 5))
                block.append(flow)
                block.append(Spacer(1, 4))

        # 제시문(보기·안내문 등) — 선지 위에 번호 없이 그대로 표시.
        # 별도 상자로 감싸지 않는다. (선지인지 보기인지 추측해서 상자로 묶으면
        #  선지가 많은 문제에서 잘못 분류될 수 있어, 원문 그대로 두는 편이 안전하다)
        for text, _is_bogi in qd.get("presented", []):
            block.append(Paragraph(esc(text), note_style))

        if qd.get("choices"):
            block.append(Spacer(1, 3))
        for i, ch in enumerate(qd.get("choices", []), start=1):
            # 요청 형식: 1) 2) 3) ...  (번호는 옅은 남색으로 살짝 구분)
            block.append(Paragraph(
                f"<font color='#6B8CAD'>{i})</font>&nbsp;{esc(ch)}", choice_style))
        # 문제 사이 여백 (다음 문제의 구분선이 위 여백을 담당)
        block.append(Spacer(1, 8))
        story.append(KeepTogether(block))

    doc.build(story)
    buf.seek(0)
    return list(PdfReader(buf).pages)


# ---------------------------------------------------------------------------
# 메인 추출 파이프라인
# ---------------------------------------------------------------------------
# 같은 연도 안에서 학기 정렬 순서 (기말 -> 중간)
_TERM_ORDER = {"기말": 0, "중간": 1}


def build_pdf(folder, subject, rows, out_path, log=print):
    """
    rows: [(year, term, [qnum, ...]), ...]  (줄마다 중간/기말 지정)
    반환: (출력 PDF 경로, 추출된 문제 수, 경고 리스트)

    엑셀 없이 동작한다. 각 문제는 풀이 PDF에서 먼저 찾고,
    없으면 시험 DOCX에서 파싱해 재렌더링한다.
    """
    warnings = []

    # 선택 항목 정리
    selections = []   # (year, term, qnum)
    for year, term, qnums in rows:
        for q in qnums:
            selections.append((year, term, q))

    # 연도 최신 -> 오래된 순. 같은 연도는 기말 -> 중간, 그 안에서 문제번호 오름차순
    selections.sort(key=lambda x: (-x[0], _TERM_ORDER.get(x[1], 9), x[2]))

    # 3) 원본별 캐시
    sol_cache = {}    # (year, term) -> (blocks, reader)
    doc_cache = {}    # (year, term) -> qmap

    writer = PdfWriter()
    stats = {"added": 0}
    cur_key = None    # (year, term)
    pending = []      # 현재 구간의 재렌더링 대기 docx 문제
    pending_title = []   # 아직 내보내지 않은 표지 [(제목, 부제)]

    def emit_title():
        """표지는 그 구간에 실제로 들어갈 내용이 확인된 뒤에만 내보낸다.

        구간 시작 시점에 미리 넣으면, 해당 연도의 파일이 없어 아무 문제도
        추출되지 않았을 때 내용 없는 표지만 덩그러니 남는다.
        """
        if pending_title:
            title, subtitle = pending_title.pop()
            writer.add_page(make_title_page(title, subtitle))

    def flush_pending():
        if not pending:
            return
        y, t = cur_key
        pages = render_docx_questions_to_pages(
            list(pending), f"{y} {subject} {t}")
        emit_title()
        for p in pages:
            writer.add_page(p)
        stats["added"] += len(pending)
        pending.clear()

    for (year, term, q) in selections:
        key = (year, term)
        # 새로운 (연도, 학기) 구간이 시작되면: 이전 대기분 출력 -> 표지 삽입
        if key != cur_key:
            flush_pending()
            cur_key = key
            pending_title[:] = [(f"{year}", f"{subject} {term}")]
        placed = False

        # 1순위: 풀이 PDF에서 (문제+해설) 블록 추출
        src = find_source(folder, year, subject, term, "풀이")
        if src:
            if key not in sol_cache:
                log(f"풀이 PDF 분석: {os.path.basename(src)}")
                sol_cache[key] = map_solution_blocks(src)
            blocks, reader = sol_cache[key]
            if q in blocks:
                flush_pending()   # 순서 유지를 위해 대기 중 docx 먼저 출력
                emit_title()
                s, e = blocks[q]
                slides = [reader.pages[p] for p in range(s, e)]
                # 문제+해설을 한 페이지에 모아 배치 (4장 초과면 4장씩 나눠 여러 페이지)
                composed = compose_slides_grid(
                    slides, f"{year} {subject} {term}  {q}번", per_page=4)
                for cp in composed:
                    writer.add_page(cp)
                stats["added"] += 1
                placed = True
                log(f"  ✓ {year} {term} {q}번 (풀이 {len(slides)}장 → {len(composed)}p)")

        # 2순위: 시험 DOCX에서 파싱해 재렌더링
        if not placed:
            docx = find_source(folder, year, subject, term, "문제")
            if docx:
                if key not in doc_cache:
                    log(f"시험 DOCX 파싱: {os.path.basename(docx)}")
                    doc_cache[key] = parse_docx_questions(docx)
                qmap = doc_cache[key]
                if q in qmap:
                    pending.append(qmap[q])
                    placed = True
                    log(f"  ✓ {year} {term} {q}번 (문제 재렌더링)")
                else:
                    warnings.append(
                        f"{year} {term} {q}번: 풀이 PDF와 시험 DOCX 어디에서도 찾지 못함")
            elif src:
                warnings.append(
                    f"{year} {term} {q}번: 풀이 PDF에 블록이 없고, 시험 DOCX 파일도 없음")
            else:
                warnings.append(f"{year} {term}: 풀이 PDF와 시험 DOCX 파일이 모두 없음")

    flush_pending()   # 마지막 구간 대기분 출력
    added = stats["added"]

    if added == 0:
        raise RuntimeError("추출된 문제가 없습니다. 파일 이름/입력을 확인하세요.")

    with open(out_path, "wb") as f:
        writer.write(f)
    log(f"완료: {added}개 문제 → {os.path.basename(out_path)}")
    return out_path, added, warnings
